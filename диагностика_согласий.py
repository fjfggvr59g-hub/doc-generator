import pandas as pd
from docx import Document
import re

EXCEL_FILE = 'Данные для согласий.xlsx'
TEMPLATE_FILE = 'Шаблон_Согласия.docx'

print("="*60)
print("🔍 ДИАГНОСТИКА: ПОИСК ПРОБЛЕМЫ")
print("="*60)

# 1. Загружаем Excel
print("\n1️⃣ ЗАГРУЖАЮ ДАННЫЕ ИЗ EXCEL...")
data = pd.read_excel(EXCEL_FILE)
print(f"   ✅ Найдено записей: {len(data)}")
print(f"   📊 Названия колонок: {list(data.columns)}")

# Показываем данные первого человека для проверки
first_person = data.iloc[0]
print("\n   👤 Данные первого человека для проверки:")
for col in data.columns:
    print(f"      {col}: {first_person[col]}")

# 2. Проверяем шаблон Word
print("\n2️⃣ ПРОВЕРЯЮ ШАБЛОН WORD...")
doc = Document(TEMPLATE_FILE)

# Ищем все метки в документе
found_placeholders = set()

print("\n   📝 Проверяю все параграфы и таблицы:")
for idx, paragraph in enumerate(doc.paragraphs):
    text = paragraph.text
    matches = re.findall(r'\{\{(.*?)\}\}', text)
    if matches:
        print(f"\n   🔹 Параграф {idx+1}:")
        print(f"      Текст: {text[:100]}...")  # Показываем начало текста
        for m in matches:
            print(f"        Найдена метка: {{ {m} }}")
            found_placeholders.add(m.strip())

# Проверяем таблицы
for table_idx, table in enumerate(doc.tables):
    print(f"\n   🔹 Таблица #{table_idx+1}:")
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            text = cell.text
            matches = re.findall(r'\{\{(.*?)\}\}', text)
            if matches:
                print(f"      Ячейка ({row_idx+1}, {cell_idx+1}):")
                print(f"        Текст: {text[:100]}...")
                for m in matches:
                    print(f"          Найдена метка: {{ {m} }}")
                    found_placeholders.add(m.strip())

# 3. Сравниваем метки из Word с колонками из Excel
print("\n" + "="*60)
print("3️⃣ СРАВНЕНИЕ МЕТОК ИЗ WORD И КОЛОНОК ИЗ EXCEL")
print("="*60)

excel_cols = set(data.columns)
word_placeholders = found_placeholders

print(f"\n📌 Метки в Word, которых НЕТ в Excel:")
missing = word_placeholders - excel_cols
for m in missing:
    print(f"   ❌ '{m}' — нет в Excel")

print(f"\n📌 Колонки в Excel, которые НЕ используются в Word:")
unused = excel_cols - word_placeholders
for col in unused:
    print(f"   ⚠️ '{col}' — есть в Excel, но нет метки в Word")

print(f"\n📌 Полное совпадение (✅):")
match = word_placeholders & excel_cols
for m in match:
    print(f"   ✅ '{m}' — есть и там, и там")

print("\n" + "="*60)
input("\nНажмите Enter для выхода...")