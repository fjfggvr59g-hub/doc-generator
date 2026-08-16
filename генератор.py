import pandas as pd
from docx import Document
import os
import re
from num2words import num2words

EXCEL_FILE = 'сотрудники.xlsx'
TEMPLATE_FILE = 'шаблон.docx'
OUTPUT_FOLDER = 'Готовые договоры'

os.makedirs(OUTPUT_FOLDER, exist_ok=True)

print("📂 Загружаю данные...")
data = pd.read_excel(EXCEL_FILE)
print(f"✅ Найдено сотрудников: {len(data)}")

def replace_all_placeholders(doc, replacements):
    """Заменяет все метки во всем документе (включая таблицы и колонтитулы)"""
    # Заменяем в параграфах
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            placeholder = f'{{{{{key}}}}}'
            if placeholder in paragraph.text:
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(value))
    
    # Заменяем в таблицах
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        placeholder = f'{{{{{key}}}}}'
                        if placeholder in paragraph.text:
                            for run in paragraph.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, str(value))
    
    # Заменяем в колонтитулах
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            for key, value in replacements.items():
                placeholder = f'{{{{{key}}}}}'
                if placeholder in paragraph.text:
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(value))
        for paragraph in section.footer.paragraphs:
            for key, value in replacements.items():
                placeholder = f'{{{{{key}}}}}'
                if placeholder in paragraph.text:
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(value))

for index, row in data.iterrows():
    doc = Document(TEMPLATE_FILE)
    
    # Создаем словарь замен
    replacements = {}
    for col in data.columns:
        replacements[col] = str(row[col]) if pd.notna(row[col]) else ''
    
    # Добавляем сумму прописью
    if 'Сумма' in data.columns and pd.notna(row['Сумма']):
        try:
            num_value = float(str(row['Сумма']).replace(',', '.'))
            sum_text = num2words(num_value, lang='ru', to='currency', currency='RUB')
            replacements['Сумма прописью'] = sum_text.capitalize()
        except:
            replacements['Сумма прописью'] = str(row['Сумма'])
    
    # Принудительно заменяем все метки
    replace_all_placeholders(doc, replacements)
    
    # Сохраняем
    fio = row.get('ФИО', f'Сотрудник_{index+1}')
    doc.save(os.path.join(OUTPUT_FOLDER, f"Договор_{fio}.docx"))
    print(f"✅ Создан договор для: {fio}")

print(f"\n🎉 Готово! Создано {len(data)} договоров")
input("Нажми Enter для выхода...")