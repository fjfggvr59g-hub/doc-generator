import pandas as pd
from docx import Document
import os
import re
from datetime import datetime

# ===== НАСТРОЙКИ =====
EXCEL_FILE = 'Данные.xlsx'
TEMPLATE_FILE = 'Шаблон_Согласия.docx'
OUTPUT_FOLDER = 'согласия'
# =====================

os.makedirs(OUTPUT_FOLDER, exist_ok=True)

print("📂 Загружаю данные...")
data = pd.read_excel(EXCEL_FILE)

# Очищаем названия колонок
data.columns = data.columns.str.strip()

print(f"✅ Найдено записей: {len(data)}")

# Словарь для перевода месяцев
MONTHS = {
    '01': 'января', '02': 'февраля', '03': 'марта',
    '04': 'апреля', '05': 'мая', '06': 'июня',
    '07': 'июля', '08': 'августа', '09': 'сентября',
    '10': 'октября', '11': 'ноября', '12': 'декабря'
}

def format_date(value):
    if pd.isna(value) or not str(value).strip():
        return ''
    
    date_str = str(value).strip()
    
    try:
        if isinstance(value, (pd.Timestamp, datetime)):
            dt = value
        elif re.match(r'\d{2}\.\d{2}\.\d{4}', date_str):
            day, month, year = date_str.split('.')
            dt = datetime(int(year), int(month), int(day))
        elif re.match(r'\d{4}-\d{2}-\d{2}', date_str):
            year, month, day = date_str.split('-')
            dt = datetime(int(year), int(month), int(day))
        else:
            return date_str
    except Exception:
        return date_str
    
    day = str(dt.day)
    month = MONTHS.get(f'{dt.month:02d}', '')
    year = str(dt.year)
    
    return f'{day} {month} {year} г.'

def format_date_with_quotes(value):
    formatted = format_date(value)
    if not formatted:
        return '«______» _______________ 20____ г.'
    parts = formatted.split(' ', 1)
    if len(parts) == 2:
        return f'«{parts[0]}» {parts[1]}'
    return formatted

def clean_text(value):
    if pd.isna(value):
        return ''
    return str(value).strip()

def format_fio_initials(full_name):
    if pd.isna(full_name) or not full_name:
        return ''
    parts = str(full_name).strip().split()
    if len(parts) == 0:
        return ''
    surname = parts[0]
    initials = '.'.join([part[0] for part in parts[1:]]) if len(parts) > 1 else ''
    if initials:
        return f"{surname} {initials}."
    return surname

def replace_in_paragraph_safe(paragraph, replacements):
    full_text = paragraph.text
    
    has_placeholder = False
    for key in replacements.keys():
        if f'{{{{{key}}}}}' in full_text:
            has_placeholder = True
            break
    
    if not has_placeholder:
        return
    
    for key, value in replacements.items():
        placeholder = f'{{{{{key}}}}}'
        if placeholder in full_text:
            full_text = full_text.replace(placeholder, str(value))
    
    paragraph.clear()
    paragraph.add_run(full_text)

# Находим все метки в шаблоне
print("📝 Анализирую шаблон...")
template_doc = Document(TEMPLATE_FILE)
placeholders = set()

for paragraph in template_doc.paragraphs:
    matches = re.findall(r'\{\{(.*?)\}\}', paragraph.text)
    for m in matches:
        placeholders.add(m.strip())
for table in template_doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                matches = re.findall(r'\{\{(.*?)\}\}', paragraph.text)
                for m in matches:
                    placeholders.add(m.strip())

print(f"📝 Найдены метки: {placeholders}")

for index, row in data.iterrows():
    doc = Document(TEMPLATE_FILE)
    
    replacements = {}
    
    # Основные поля
    for col in data.columns:
        if col in placeholders:
            if col == 'Дата рождения':
                replacements[col] = format_date(row.get(col, ''))
            elif col == 'Дата согласия':
                replacements[col] = format_date_with_quotes(row.get(col, ''))
            else:
                replacements[col] = clean_text(row.get(col, ''))
    
    if 'Фамилия И.О.' in placeholders:
        replacements['Фамилия И.О.'] = format_fio_initials(row.get('ФИО', ''))
    
    for paragraph in doc.paragraphs:
        replace_in_paragraph_safe(paragraph, replacements)
    
    for table in doc.tables:
        for row_table in table.rows:
            for cell in row_table.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph_safe(paragraph, replacements)
    
    fio = clean_text(row.get('ФИО', f'Согласие_{index+1}'))
    filename = f"Согласие_{fio}.docx"
    filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
    doc.save(os.path.join(OUTPUT_FOLDER, filename))
    print(f"✅ Создано: {filename}")

print(f"\n🎉 Готово! Создано {len(data)} согласий")
